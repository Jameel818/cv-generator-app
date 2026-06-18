from flask import Flask, render_template, request, send_file, jsonify, make_response
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib import colors  # ← السطر الجديد 1
import io
import os
import arabic_reshaper
from bidi.algorithm import get_display
app = Flask(__name__)

# === هنا سجلنا خط Cairo ===
FONT_PATH = os.path.join(os.path.dirname(__file__), 'Cairo-Regular.ttf')
pdfmetrics.registerFont(TTFont('Cairo', FONT_PATH))

# الخطوط المتاحة
FONTS = ['Arial', 'Times New Roman', 'Cairo']
SIZES = [12, 14, 16, 18, 20, 24]

@app.route('/')
def landing():
    return render_template('landing.html')  # صفحة تسويقية جديدة

@app.route('/app')  # ← هذا الجديد
def index():
    return render_template('index.html', fonts=FONTS, sizes=SIZES)  # تطبيقك زي ما هو

@app.route('/generate_word', methods=['POST'])
def generate_word():
    data = request.json
    doc = Document()

    # العنوان
    title = doc.add_paragraph()
    title_run = title.add_run(data['title'])
    title_run.font.name = data['font_family']
    title_run.font.size = Pt(data['font_size'])
    title_run.font.color.rgb = RGBColor.from_string(data['font_color'].replace('#', ''))
    title_run.bold = True

    # المحتوى
    content = doc.add_paragraph()
    content_run = content.add_run(data['content'])
    content_run.font.name = data['font_family']
    content_run.font.size = Pt(data['font_size'])
    content_run.font.color.rgb = RGBColor.from_string(data['font_color'].replace('#', ''))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name='document.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    data = request.json
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    styles = getSampleStyleSheet()
    
    styles.add(ParagraphStyle(
        name='Arabic', 
        fontName='Cairo',
        fontSize=data['font_size'],
        alignment=TA_RIGHT, 
        leading=24,
        textColor=colors.HexColor(data['font_color'])  # ← السطر الجديد 2: هذا هو الحل
    ))

    def reshape_arabic(text):
        reshaped_text = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped_text)
        bidi_text = bidi_text.replace('‌', ' ').replace('‍', '')  # نحذف حروف التشكيل الزايدة
        return bidi_text

    story = []
    title_text = reshape_arabic(data['title'])
    content_text = reshape_arabic(data['content']).replace('\n', '<br/>')
    
    title = Paragraph(f"<b>{title_text}</b>", styles['Arabic'])
    content = Paragraph(content_text, styles['Arabic'])

    story.append(title)
    story.append(Spacer(1, 12))
    story.append(content)

    doc.build(story)
    
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    response = make_response(pdf_bytes)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'attachment; filename=document.pdf'
    response.headers['Content-Length'] = len(pdf_bytes)
    response.headers['Accept-Ranges'] = 'none'
    
    return response

if __name__ == '__main__':
    app.run(debug=False)